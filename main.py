#!/usr/bin/env python3
"""
AI-Powered Cover Letter Generator - Main Entry Point
"""

import argparse
import sys
import os
from cover_letter_generator import CoverLetterGenerator
from datetime import datetime


def create_pdf_cover_letter(cover_letter_text, output_filename="coverLetter.pdf"):
    """Create a beautiful PDF version of the cover letter."""
    try:
        from reportlab.lib.pagesizes import letter
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import inch
        from reportlab.lib.colors import HexColor
        
        # Create PDF document with reduced top margin
        doc = SimpleDocTemplate(
            output_filename,
            pagesize=letter,
            topMargin=0.5 * inch,
            bottomMargin=0.75 * inch,
            leftMargin=1.0 * inch,
            rightMargin=1.0 * inch,
        )
        story = []
        
        # Define styles
        styles = getSampleStyleSheet()
        
        # Custom title style
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=17,
            spaceAfter=10,
            textColor=HexColor('#000000'),
            alignment=1,  # Center alignment
            fontName='Times-Bold'
        )
        
        # Custom body style
        body_style = ParagraphStyle(
            'CustomBody',
            parent=styles['Normal'],
            fontSize=12,
            spaceAfter=12,
            leading=16,
            fontName='Times-Roman'
        )
        
        # Add title
        story.append(Paragraph("Cover Letter", title_style))
        story.append(Spacer(1, 6))
        
        # Split cover letter into paragraphs and add them
        paragraphs = cover_letter_text.split('\n\n')
        for para in paragraphs:
            if para.strip():
                # Check if this paragraph contains contact information (multiple lines)
                if '\n' in para and any(keyword in para.lower() for keyword in ['email:', 'phone:', 'linkedin:']):
                    # Split contact info into separate lines and add each as a separate paragraph
                    contact_lines = para.strip().split('\n')
                    for line in contact_lines:
                        if line.strip():
                            story.append(Paragraph(line.strip(), body_style))
                            story.append(Spacer(1, 6))  # Less spacing between contact lines
                else:
                    # Regular paragraph
                    story.append(Paragraph(para.strip(), body_style))
                    story.append(Spacer(1, 12))
        
        # Build PDF
        doc.build(story)
        return True
        
    except ImportError:
        print("Warning: reportlab not installed. Install with: pip install reportlab")
        return False
    except Exception as e:
        print(f"Error creating PDF: {e}")
        return False


def create_pages_document(cover_letter_text, output_filename="coverLetter.docx"):
    """Create a Pages-compatible document (.docx format)."""
    try:
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        # Create document
        doc = Document()
        
        # Set margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(0.8)
            section.bottom_margin = Inches(0.8)
            section.left_margin = Inches(1.0)
            section.right_margin = Inches(1.0)
        
        # Add title as plain paragraph (no heading style), black, centered
        title_p = doc.add_paragraph()
        title_run = title_p.add_run('Cover Letter')
        title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run.font.name = 'Times New Roman'
        title_run.font.size = Pt(17)
        title_run.font.bold = True
        title_run.font.color.rgb = RGBColor(0, 0, 0)
        
        # Add minimal spacing after title
        doc.add_paragraph()
        
        # Split cover letter into paragraphs and add them
        paragraphs = cover_letter_text.split('\n\n')
        for para in paragraphs:
            if para.strip():
                p = doc.add_paragraph(para.strip())
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                
                # Style the paragraph
                for run in p.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(12)
        
        # Save document
        doc.save(output_filename)
        return True
        
    except ImportError:
        print("Warning: python-docx not installed. Install with: pip install python-docx")
        return False
    except Exception as e:
        print(f"Error creating Pages document: {e}")
        return False


def create_timestamped_directory():
    """Create a timestamped directory for this run's output files."""
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    output_dir = os.path.join("coverletters", f"cover_letter_{timestamp}")
    os.makedirs(output_dir, exist_ok=True)
    return output_dir


def main():
    """Main function to handle command line arguments and generate cover letter."""
    parser = argparse.ArgumentParser(
        description="Generate a cover letter based on a resume and job description using GPT-4o"
    )
    parser.add_argument(
        "resume",
        nargs='?',
        default="ArtinMajd_CV.pdf",
        help="Resume text or path to file containing resume (default: Artin_Majd_CV.pdf)"
    )
    parser.add_argument(
        "job_description",
        nargs='?',
        default="job_description.txt",
        help="Job description text or path to file containing job description (default: job_description.txt)"
    )
    parser.add_argument(
        "--output", "-o",
        help="Output file path (default: cover_letter.txt)"
    )
    parser.add_argument(
        "--model", "-m",
        default="gpt-4o",
        help="OpenAI model to use (default: gpt-4o)"
    )
    parser.add_argument(
        "--no-pdf",
        action="store_true",
        help="Skip PDF generation"
    )
    parser.add_argument(
        "--use-knowledge-base", "--kb",
        action="store_true",
        help="Use knowledgeBase.txt instead of the resume"
    )
    
    args = parser.parse_args()
    
    try:
        # Initialize the cover letter generator
        generator = CoverLetterGenerator()
        
        # Create timestamped output directory
        output_dir = create_timestamped_directory()
        print(f"📁 Output directory created: {output_dir}")
        
        # If requested, override resume with knowledgeBase.txt
        if args.use_knowledge_base:
            kb_path = "knowledgeBase.txt"
            if not os.path.isfile(kb_path):
                raise FileNotFoundError(f"Knowledge base file not found: {kb_path}")
            effective_resume = kb_path
        else:
            effective_resume = args.resume
        
        # Generate the cover letter
        cover_letter = generator.generate_with_ai(
            resume=effective_resume,
            job_description=args.job_description,
            model=args.model
        )
        
        # Output the result
        if args.output:
            # If custom output name provided, use it in the timestamped directory
            txt_filename = os.path.join(output_dir, args.output)
            with open(txt_filename, 'w') as f:
                f.write(cover_letter)
            print(f"Cover letter saved to {txt_filename}")
        else:
            # Create default output file in timestamped directory
            txt_filename = os.path.join(output_dir, "coverLetter.txt")
            with open(txt_filename, 'w') as f:
                f.write(cover_letter)
            print(f"Cover letter saved to {txt_filename}")
        
        # Generate PDF version
        if not args.no_pdf:
            print("\n" + "="*50)
            print("GENERATING PDF VERSION")
            print("="*50)
            
            if args.output:
                pdf_filename = os.path.join(output_dir, args.output.replace('.txt', '.pdf').replace('.doc', '.pdf'))
            else:
                pdf_filename = os.path.join(output_dir, "coverLetter.pdf")
            
            if create_pdf_cover_letter(cover_letter, pdf_filename):
                print(f"✅ Beautiful PDF created: {pdf_filename}")
            else:
                print("❌ PDF generation failed")
        
        # Generate Pages document
        print("\n" + "="*50)
        print("GENERATING PAGES DOCUMENT")
        print("="*50)
        
        if args.output:
            pages_filename = os.path.join(output_dir, args.output.replace('.txt', '.docx').replace('.pdf', '.docx'))
        else:
            pages_filename = os.path.join(output_dir, "coverLetter.docx")
        
        if create_pages_document(cover_letter, pages_filename):
            print(f"✅ Pages document created: {pages_filename}")
            print("   (Open with Pages app or any Word-compatible application)")
        else:
            print("❌ Pages document generation failed")
        
        print(f"\n🎉 All files saved to: {output_dir}")
        
    except Exception as e:
        print(f"Error: {e}", file=sys.stderr)
        sys.exit(1)


if __name__ == "__main__":
    main()
